"""
01_documents.py — STAGE 1 of the RAG pipeline:  **Raw Documents**

    >>> Raw Documents <<<  →  Preprocessing  →  Chunking  →  Vector Representation
    →  Vector Store  →  Context Retrieval  →  Prompt Building  →  LLM Generation  →  Streamlit UI

# WHY does this stage exist?
# Everything downstream is only as good as what we feed it. Retrieval cannot find
# a fact that was never ingested, and the LLM is forbidden (by the prompt in stage
# 07) from using anything except what retrieval returns. So this file has exactly
# one job: turn *any* source of raw text into ONE uniform `Document` shape, so
# that stages 02-07 never have to care where the text came from.
#
# Two very different sources have to arrive at the same shape:
#   1. The built-in pharmaceutical corpus — 506 real FDA drug labels, structured
#      as JSON records with named fields (mechanism, dosage, side_effects, …).
#   2. Files the user uploads in the Streamlit app — PDF / DOCX / TXT / MD / CSV /
#      JSON, which are unstructured blobs with no fields at all.
#
# The `Document` dataclass below is that common shape. Structured drug records
# explode into one Document per field (so a citation can say "field=dosage");
# uploaded files become one Document per page/section (so a citation can say
# "page 3"). Either way stage 03 receives a plain list of Documents.

Run standalone:
    python 01_documents.py                 # corpus stats + a sample document
    python 01_documents.py --refresh 500   # re-download drug labels from openFDA
"""
from __future__ import annotations

import io
import json
import re
import urllib.request
from dataclasses import dataclass, field as dc_field
from pathlib import Path
from typing import Iterable, Optional

# ─────────────────────────────────────────────────────────────────────────────
# Paths & configuration
# ─────────────────────────────────────────────────────────────────────────────
HERE = Path(__file__).resolve().parent
DATA_DIR = HERE / "data"

# Two shipped corpora. `drugs_large.json` (506 real openFDA labels) is the
# production corpus; `drugs.json` (12 hand-written drugs) is the small teaching
# corpus the Learning-Edition notebook uses. We prefer the large one and fall
# back, so the project still runs if the big file is missing.
DRUGS_LARGE = DATA_DIR / "drugs_large.json"
DRUGS_SMALL = DATA_DIR / "drugs.json"

# WHY these seven fields, in this order?
# They are the fields an FDA Structured Product Label actually provides, and the
# order runs general → specific → safety. Retrieval metadata uses these names
# verbatim, so a citation reads "metformin · dosage" instead of "chunk #4127".
FIELD_LABELS = ["description", "mechanism", "indications", "dosage",
                "side_effects", "contraindications", "interactions"]

# Upload limits — a defensive cap so a 500 MB PDF cannot take the app down.
MAX_UPLOAD_CHARS = 400_000


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1
# Define the single uniform document shape used by the whole pipeline.
#
# WHY a dataclass and not a bare dict?
# Because every downstream stage reads `.title`, `.field` and `.text` by name.
# A dataclass makes that contract explicit and gives us a typo-proof constructor;
# a dict would let stage 03 silently read a key that stage 01 never wrote.
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class Document:
    """One retrievable unit of raw text, before chunking.

    doc_id   : stable unique id, e.g. "metformin::dosage" or "notes.pdf::p3"
    title    : human-facing name shown in citations ("Metformin Hydrochloride")
    field    : which part of the source this is ("dosage", "page 3", "body")
    category : coarse grouping, used for filtering ("biguanide", "upload")
    text     : the raw text itself — NOT yet cleaned, NOT yet chunked
    origin   : provenance string, e.g. "drug-corpus" or "upload:notes.pdf"
    """
    doc_id: str
    title: str
    field: str
    category: str
    text: str
    origin: str = "drug-corpus"
    meta: dict = dc_field(default_factory=dict)

    def __post_init__(self):
        # Normalize whitespace at the door. Every later stage may now assume
        # `text` is a single-spaced string, which keeps the word-based chunker in
        # stage 03 honest (it counts words with .split()).
        self.text = re.sub(r"\s+", " ", str(self.text or "")).strip()

    @property
    def word_count(self) -> int:
        return len(self.text.split())


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2
# Load the built-in structured drug corpus from JSON.
#
# WHY keep the raw records around instead of going straight to Documents?
# Because two other features need the *record* view, not the document view:
# the entity-grounding vocabulary in stage 06 (which drugs do we actually cover?)
# and the sidebar drug count in the UI. So we expose both.
# ─────────────────────────────────────────────────────────────────────────────
def default_corpus_path() -> Path:
    """Prefer the 506-drug openFDA corpus; fall back to the 12-drug teaching set."""
    return DRUGS_LARGE if DRUGS_LARGE.exists() else DRUGS_SMALL


def load_drug_records(path: Optional[str | Path] = None) -> list[dict]:
    """Read the drug JSON file and return the raw list of drug records."""
    p = Path(path) if path else default_corpus_path()
    if not p.exists():
        raise FileNotFoundError(
            f"Drug corpus not found at {p}. Expected data/drugs_large.json or "
            f"data/drugs.json — run `python 01_documents.py --refresh 500` to rebuild it."
        )
    records = json.loads(p.read_text(encoding="utf-8"))
    if not isinstance(records, list) or not records:
        raise ValueError(f"{p} did not contain a non-empty JSON list of drug records.")
    return records


def documents_from_drug_records(records: Iterable[dict]) -> list[Document]:
    """Explode each structured drug record into one Document per populated field.

    # STEP 2a
    # WHY one Document per field instead of one per drug?
    # A whole drug label is ~3000 words spanning mechanism, dosing, warnings and
    # interactions. Embedding that as a single vector produces a blurry average
    # that matches every drug question weakly and none of them strongly. Splitting
    # on the label's own field boundaries is free, semantically clean chunking:
    # each Document already contains exactly one coherent topic.
    """
    docs: list[Document] = []
    for rec in records:
        name = str(rec.get("name", "")).strip()
        if not name:
            continue
        category = str(rec.get("category", "") or "uncategorized").strip()
        for fld in FIELD_LABELS:
            body = rec.get(fld, "")
            if not body:
                continue  # skip empty fields — an empty chunk is pure noise in the index
            docs.append(Document(
                doc_id=f"{name.lower()}::{fld}",
                title=name,
                field=fld,
                category=category,
                text=str(body),
                origin="drug-corpus",
                meta={"source": rec.get("source", "")},
            ))
    return docs


def load_drug_documents(path: Optional[str | Path] = None) -> list[Document]:
    """Convenience: read the JSON corpus and return it as Documents."""
    return documents_from_drug_records(load_drug_records(path))


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3
# Ingest user-uploaded files (the Streamlit "Upload documents" feature).
#
# WHY does an upload path belong in stage 01 rather than in the UI?
# Because the UI must not own pipeline logic. Streamlit hands us bytes + a
# filename; this function turns those into the same `Document` objects the drug
# corpus produces, so an uploaded PDF flows through chunking → embedding →
# retrieval → citation with zero special-casing anywhere downstream.
#
# Every parser is optional and degrades gracefully: if `pypdf` is missing we say
# so clearly instead of crashing the app.
# ─────────────────────────────────────────────────────────────────────────────
def _read_pdf(data: bytes, filename: str) -> list[tuple[str, str]]:
    """Return [(section_label, text), …] for a PDF — one entry per page.

    WHY per page? Page numbers are the only citation anchor a PDF reliably has,
    so keeping the page boundary lets the UI say "notes.pdf · page 4".
    """
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover - depends on optional dep
        raise RuntimeError(
            "Reading PDFs requires the `pypdf` package (`pip install pypdf`)."
        ) from exc
    reader = PdfReader(io.BytesIO(data))
    out = []
    for i, page in enumerate(reader.pages, 1):
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            out.append((f"page {i}", text))
    return out


def _read_docx(data: bytes, filename: str) -> list[tuple[str, str]]:
    """Return [(section_label, text)] for a .docx — paragraphs joined into one body."""
    try:
        import docx  # python-docx
    except Exception as exc:  # pragma: no cover - depends on optional dep
        raise RuntimeError(
            "Reading .docx requires the `python-docx` package (`pip install python-docx`)."
        ) from exc
    document = docx.Document(io.BytesIO(data))
    text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
    return [("body", text)] if text.strip() else []


def _read_csv(data: bytes, filename: str) -> list[tuple[str, str]]:
    """Return one entry per CSV row, rendered as `col: value` lines.

    WHY render rows as text instead of keeping them tabular?
    Retrieval works on text. Flattening a row into "drug: X | level: Major"
    makes it embeddable while keeping every column name visible to the LLM.
    """
    import csv as _csv
    text = data.decode("utf-8", errors="replace")
    rows = list(_csv.DictReader(io.StringIO(text)))
    out = []
    for i, row in enumerate(rows, 1):
        rendered = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
        if rendered.strip():
            out.append((f"row {i}", rendered))
    return out


def _read_json(data: bytes, filename: str) -> list[tuple[str, str]]:
    """Return entries for a JSON file.

    Handles the two shapes that actually turn up: a list of records (one entry
    each) and a single object (one entry). Anything else is stringified whole.
    """
    payload = json.loads(data.decode("utf-8", errors="replace"))
    if isinstance(payload, list):
        out = []
        for i, item in enumerate(payload, 1):
            if isinstance(item, dict):
                rendered = " | ".join(f"{k}: {v}" for k, v in item.items()
                                      if isinstance(v, (str, int, float)) and str(v).strip())
            else:
                rendered = str(item)
            if rendered.strip():
                out.append((f"record {i}", rendered))
        return out
    if isinstance(payload, dict):
        rendered = " | ".join(f"{k}: {v}" for k, v in payload.items()
                              if isinstance(v, (str, int, float)))
        return [("body", rendered)] if rendered.strip() else []
    return [("body", str(payload))]


def _read_plaintext(data: bytes, filename: str) -> list[tuple[str, str]]:
    """Return [(section, text)] for .txt / .md, split on markdown headings.

    WHY split on headings? A markdown heading is an author-declared topic
    boundary — the cheapest possible semantic chunk hint. If there are no
    headings we return one body section and let stage 03's sliding window
    handle the split instead.
    """
    text = data.decode("utf-8", errors="replace")
    # Split before lines that look like "# Heading" / "## Heading".
    parts = re.split(r"\n(?=#{1,6}\s+)", text)
    if len(parts) == 1:
        return [("body", text)] if text.strip() else []
    out = []
    for part in parts:
        if not part.strip():
            continue
        head = part.strip().splitlines()[0].lstrip("# ").strip()
        out.append((head[:60] or "section", part))
    return out


# Dispatch table: file suffix → parser. Adding a format means adding one line.
_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".csv": _read_csv,
    ".json": _read_json,
    ".txt": _read_plaintext,
    ".md": _read_plaintext,
    ".markdown": _read_plaintext,
}

SUPPORTED_UPLOAD_TYPES = sorted(s.lstrip(".") for s in _READERS)


def documents_from_upload(filename: str, data: bytes,
                          category: str = "upload") -> list[Document]:
    """Turn one uploaded file into Documents. Raises RuntimeError with a readable
    message if the format is unsupported or its parser dependency is missing."""
    suffix = Path(filename).suffix.lower()
    reader = _READERS.get(suffix)
    if reader is None:
        raise RuntimeError(
            f"Unsupported file type '{suffix or filename}'. "
            f"Supported: {', '.join(SUPPORTED_UPLOAD_TYPES)}."
        )
    sections = reader(data, filename)
    stem = Path(filename).stem

    docs: list[Document] = []
    budget = MAX_UPLOAD_CHARS
    for label, text in sections:
        if budget <= 0:
            break
        text = text[:budget]
        budget -= len(text)
        doc = Document(
            doc_id=f"{filename}::{label}",
            title=stem,                 # citations read "notes · page 3"
            field=label,
            category=category,
            text=text,
            origin=f"upload:{filename}",
            meta={"filename": filename},
        )
        if doc.text:                    # __post_init__ may have emptied it
            docs.append(doc)
    if not docs:
        raise RuntimeError(f"No readable text found in '{filename}'.")
    return docs


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4
# (Optional) Rebuild the drug corpus live from the openFDA API.
#
# WHY keep the downloader in the repo at all, since the JSON is committed?
# Reproducibility. The committed corpus is a cache, and this function is the
# recipe that produced it. `--refresh` regenerates it from the authoritative
# source, so the dataset is never a mystery blob.
# ─────────────────────────────────────────────────────────────────────────────
_OPENFDA_SEARCH = ("_exists_:indications_and_usage+AND+_exists_:mechanism_of_action"
                   "+AND+_exists_:dosage_and_administration")

# our schema field  ->  openFDA label field
_OPENFDA_FIELD_MAP = {
    "indications": "indications_and_usage",
    "mechanism": "mechanism_of_action",
    "dosage": "dosage_and_administration",
    "side_effects": "adverse_reactions",
    "contraindications": "contraindications",
    "interactions": "drug_interactions",
    "description": "description",
}
_MAX_FIELD_CHARS = 4000  # keep one field from dwarfing the whole record


def _clean_fda_value(value) -> str:
    if isinstance(value, list):
        value = " ".join(str(x) for x in value)
    return re.sub(r"\s+", " ", str(value)).strip()[:_MAX_FIELD_CHARS]


def download_openfda(target: int = 500, per_page: int = 100, verbose: bool = True) -> list[dict]:
    """Page through the public openFDA drug-label API and build drug records.

    No API key is needed. We keep only labels rich enough to be useful (they must
    have indications + mechanism + at least 4 populated fields) and deduplicate by
    generic name, because the same molecule appears under dozens of manufacturers.
    """
    seen, drugs, skip = set(), [], 0
    while len(drugs) < target and skip < 25_000:
        url = (f"https://api.fda.gov/drug/label.json?search={_OPENFDA_SEARCH}"
               f"&limit={per_page}&skip={skip}")
        try:
            with urllib.request.urlopen(url, timeout=30) as resp:
                batch = json.load(resp).get("results", [])
        except Exception as exc:
            if verbose:
                print(f"  stopped at skip={skip}: {type(exc).__name__}")
            break
        if not batch:
            break
        for rec in batch:
            openfda = rec.get("openfda", {})
            name = (openfda.get("generic_name") or openfda.get("brand_name") or [""])[0].strip().title()
            if not name or name.lower() in seen:
                continue
            if not rec.get("indications_and_usage") or not rec.get("mechanism_of_action"):
                continue
            category = (openfda.get("pharm_class_epc") or openfda.get("pharm_class_moa")
                        or openfda.get("route") or ["Unclassified"])[0].title()
            drug = {"name": name, "category": category}
            for ours, theirs in _OPENFDA_FIELD_MAP.items():
                drug[ours] = _clean_fda_value(rec.get(theirs, ""))
            if not drug["description"]:
                drug["description"] = drug["indications"][:600]
            drug["source"] = "openFDA drug label (SPL / DailyMed)"
            if sum(1 for f in _OPENFDA_FIELD_MAP if drug[f]) >= 4:
                seen.add(name.lower())
                drugs.append(drug)
        skip += per_page
        if verbose:
            print(f"  skip={skip:>5}  collected={len(drugs)}")
    return drugs


def refresh_corpus(target: int = 500, out_path: Optional[Path] = None) -> Path:
    """Download and overwrite the cached drug corpus. Returns the written path."""
    out_path = Path(out_path) if out_path else DRUGS_LARGE
    out_path.parent.mkdir(parents=True, exist_ok=True)
    drugs = download_openfda(target=target)
    if not drugs:
        raise RuntimeError("openFDA returned no usable records — corpus not overwritten.")
    out_path.write_text(json.dumps(drugs, ensure_ascii=False, indent=1), encoding="utf-8")
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# STEP 5
# Reporting helper — used by __main__ here and by the Streamlit sidebar.
# ─────────────────────────────────────────────────────────────────────────────
def corpus_stats(docs: list[Document]) -> dict:
    """Summarise a document set: counts, word totals, field distribution."""
    from collections import Counter
    words = [d.word_count for d in docs]
    return {
        "documents": len(docs),
        "sources": len({d.title for d in docs}),
        "total_words": sum(words),
        "mean_words": round(sum(words) / len(words), 1) if words else 0.0,
        "max_words": max(words) if words else 0,
        "by_field": dict(Counter(d.field for d in docs).most_common(10)),
    }


# ─────────────────────────────────────────────────────────────────────────────
# Standalone run — proves this stage works on its own.
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import argparse
    import sys

    try:  # Windows consoles choke on the arrows/bullets used below
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    ap = argparse.ArgumentParser(description="Stage 1 — raw document loading.")
    ap.add_argument("--refresh", type=int, metavar="N",
                    help="re-download N drug labels from openFDA and overwrite the cache")
    ap.add_argument("--path", help="explicit path to a drugs JSON file")
    args = ap.parse_args()

    if args.refresh:
        print(f"Downloading {args.refresh} drug labels from openFDA…")
        written = refresh_corpus(args.refresh)
        print(f"Wrote {written}")

    records = load_drug_records(args.path)
    docs = documents_from_drug_records(records)
    stats = corpus_stats(docs)

    print("=" * 70)
    print("  STAGE 1 — RAW DOCUMENTS")
    print("=" * 70)
    print(f"Corpus file    : {Path(args.path) if args.path else default_corpus_path()}")
    print(f"Drug records   : {len(records)}")
    print(f"Documents      : {stats['documents']}  (one per populated field)")
    print(f"Total words    : {stats['total_words']:,}")
    print(f"Words/document : mean {stats['mean_words']}, max {stats['max_words']}")
    print("\nDocuments per field:")
    for fld, n in stats["by_field"].items():
        print(f"  {fld:<20s} {n:>6,}")

    sample = docs[0]
    print(f"\nSample document — {sample.doc_id}")
    print(f"  title    : {sample.title}")
    print(f"  field    : {sample.field}")
    print(f"  category : {sample.category}")
    print(f"  text     : {sample.text[:200]}…")
    print(f"\nUpload formats supported: {', '.join(SUPPORTED_UPLOAD_TYPES)}")
